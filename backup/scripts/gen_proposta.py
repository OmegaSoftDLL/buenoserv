#!/usr/bin/env python3
"""Gera proposta comercial BUENOSERV em DOCX e PDF"""
import json, sys, os
from datetime import datetime, timedelta

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

LOGO_PATH = os.path.expanduser("~/.config/opencode/agents/logo-buenoserv.jpeg")

INFO = {
    "empresa": "BUENOSERV SERVIÇOS DE ENGENHARIA LTDA",
    "cnpj": "60.490.193/0001-38",
    "endereco": "Rua Giacomo Fior, 427 - Leme - SP",
    "diretor": "Ricardo Bueno",
    "email": "ricardo.bueno@buenoservengenharia.com",
    "regime": "Simples Nacional (LC 123/2006)",
    "banco": "BTG Pactual (208) | Ag: 0050 | CC: 2321479-4",
    "pix": "60.490.193/0001-38 (CNPJ)",
}

def add_styled_paragraph(doc, text, style='Normal', bold=False, size=11, color=None, alignment=None, space_after=6):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p

def gerar_proposta(cliente, contato, escopo, valor_mensal, prazo_meses, proposta_id, site=""):
    valor_anual = valor_mensal * prazo_meses
    data_envio = datetime.now()
    validade = data_envio + timedelta(days=30)

    doc = Document()

    # Estilos default
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Header - Tabela com logo e info
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Celula esquerda - dados
    cell_left = header_table.cell(0, 0)
    title = cell_left.paragraphs[0]
    run = title.add_run('BUENOSERV')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(26, 35, 126)
    subtitle = cell_left.add_paragraph('Engenharia & Telecomunicações')
    subtitle.runs[0].font.size = Pt(9)
    subtitle.runs[0].font.color.rgb = RGBColor(136, 136, 136)

    # Celula direita - logo
    if os.path.exists(LOGO_PATH):
        cell_right = header_table.cell(0, 1)
        cell_right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = cell_right.paragraphs[0].add_run()
        run.add_picture(LOGO_PATH, width=Cm(3.5))

    # Linha dourada
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('_' * 80)
    run.font.color.rgb = RGBColor(201, 168, 76)
    run.font.size = Pt(8)

    # Titulo
    add_styled_paragraph(doc, f'PROPOSTA TÉCNICA E COMERCIAL', bold=True, size=16,
                         color=RGBColor(26, 35, 126), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, f'{proposta_id}', bold=True, size=12,
                         color=RGBColor(201, 168, 76), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, f'{data_envio.strftime("%d/%m/%Y")}', size=9,
                         color=RGBColor(150, 150, 150), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # 1. Dados do cliente
    add_styled_paragraph(doc, '1. DADOS DO CLIENTE', bold=True, size=12,
                         color=RGBColor(26, 35, 126), space_after=4)
    dados = [
        ('Empresa', cliente),
        ('Contato', contato),
        ('Site', site or '-'),
    ]
    table = doc.add_table(rows=len(dados), cols=2)
    for i, (k, v) in enumerate(dados):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                paragraph.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # 2. Escopo
    add_styled_paragraph(doc, '2. ESCOPO TÉCNICO', bold=True, size=12,
                         color=RGBColor(26, 35, 126), space_after=4)
    add_styled_paragraph(doc, escopo, size=10, color=RGBColor(80, 80, 80))
    doc.add_paragraph()

    # 3. Investimento
    add_styled_paragraph(doc, '3. INVESTIMENTO', bold=True, size=12,
                         color=RGBColor(26, 35, 126), space_after=4)
    inv_table = doc.add_table(rows=4, cols=2)
    inv_data = [
        ('Valor mensal (R$)', f'{valor_mensal:,.2f}'),
        ('Prazo de execução', f'{prazo_meses} meses'),
        ('Valor total (R$)', f'{valor_anual:,.2f}'),
        ('Condição de pagamento', 'Conforme cronograma físico-financeiro'),
    ]
    for i, (k, v) in enumerate(inv_data):
        inv_table.cell(i, 0).text = k
        inv_table.cell(i, 1).text = v
        for cell in inv_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                paragraph.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # 4. Condições
    add_styled_paragraph(doc, '4. CONDIÇÕES GERAIS', bold=True, size=12,
                         color=RGBColor(26, 35, 126), space_after=4)
    condicoes = [
        f'Validade da proposta: {validade.strftime("%d/%m/%Y")}',
        'Regime tributário: Simples Nacional (LC 123/2006)',
        'Reajuste anual pelo IPCA',
        'Carga horária: 160 horas/mês',
        'Despesas de viagem reembolsáveis (transporte, hospedagem, alimentação)',
        'Garantia dos serviços: 12 meses após o aceite',
    ]
    for c in condicoes:
        add_styled_paragraph(doc, f'  - {c}', size=10, color=RGBColor(80, 80, 80), space_after=2)
    doc.add_paragraph()

    # 5. Dados bancários
    add_styled_paragraph(doc, '5. DADOS BANCÁRIOS', bold=True, size=12,
                         color=RGBColor(26, 35, 126), space_after=4)
    for line in [
        f'Banco: {INFO["banco"]}',
        f'PIX: {INFO["pix"]}',
        f'CNPJ: {INFO["cnpj"]}',
    ]:
        add_styled_paragraph(doc, line, size=10, color=RGBColor(80, 80, 80), space_after=2)
    doc.add_paragraph()

    # 6. Qualificação
    add_styled_paragraph(doc, '6. QUALIFICAÇÃO TÉCNICA', bold=True, size=12,
                         color=RGBColor(26, 35, 126), space_after=4)
    add_styled_paragraph(doc,
        'A BUENOSERV possui mais de 20 anos de experiência em projetos de telecomunicações, '
        'automação de subestações e sistemas de proteção, com equipe técnica qualificada e '
        'vasta experiência em concessionárias e indústrias do setor elétrico.',
        size=10, color=RGBColor(80, 80, 80))

    doc.add_paragraph()
    add_styled_paragraph(doc, 'Atenciosamente,', size=10, space_after=2)
    add_styled_paragraph(doc, 'Ricardo Bueno', bold=True, size=11,
                         color=RGBColor(26, 35, 126), space_after=0)
    add_styled_paragraph(doc, 'Diretor Técnico', size=9,
                         color=RGBColor(136, 136, 136), space_after=0)
    add_styled_paragraph(doc, INFO['email'], size=9,
                         color=RGBColor(136, 136, 136))

    # Salvar
    out_dir = '/tmp/opencode/propostas'
    os.makedirs(out_dir, exist_ok=True)
    out_path = f'{out_dir}/{proposta_id}.docx'
    doc.save(out_path)
    print(f'OK: {out_path}')
    return out_path

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Uso: gen_proposta.py <args_json>')
        print('Exemplo: gen_proposta.py \'{"cliente":"E4 ENERGIA","contato":"Giovanna","escopo":"...","valor_mensal":58240,"prazo_meses":12,"proposta_id":"BSE-01-R0"}\'')
        sys.exit(1)
    args = json.loads(sys.argv[1])
    gerar_proposta(
        args['cliente'],
        args.get('contato', ''),
        args.get('escopo', 'Consultoria técnica em telecomunicações para padronização de subestações SE 145kV'),
        args.get('valor_mensal', 0),
        args.get('prazo_meses', 12),
        args.get('proposta_id', 'BSE-00'),
        args.get('site', ''),
    )
