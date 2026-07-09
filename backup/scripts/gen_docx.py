#!/usr/bin/env python3
"""Gerador de documentos Word (.docx) — python-docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import sys, json

COR_AZUL = RGBColor(0x1A, 0x23, 0x7E)

def nova_pagina():
    return Document()

def salvar(doc, path):
    doc.save(path)
    print(f"DOCX|OK|{path}")

def add_capa(doc, titulo, subtitulo="", cliente="", data=""):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo)
    r.font.size = Pt(24)
    r.bold = True
    r.font.color.rgb = COR_AZUL
    if subtitulo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitulo)
        r.font.size = Pt(14)
    if cliente:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"Cliente: {cliente}")
        r.font.size = Pt(12)
    if data:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"Data: {data}")
        r.font.size = Pt(12)
    doc.add_page_break()

def add_secao(doc, titulo, conteudo):
    doc.add_heading(titulo, level=1)
    for par in conteudo:
        if par.startswith('##'):
            doc.add_heading(par.replace('##','').strip(), level=2)
        elif par.startswith('- '):
            doc.add_paragraph(par[2:], style='List Bullet')
        elif par.startswith('1.') or par.startswith('2.') or par.startswith('3.'):
            doc.add_paragraph(par, style='List Number')
        else:
            doc.add_paragraph(par)

def add_tabela(doc, cabecalhos, dados):
    table = doc.add_table(rows=1, cols=len(cabecalhos))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(cabecalhos):
        table.rows[0].cells[i].text = h
    for row_data in dados:
        row = table.add_row().cells
        for i, v in enumerate(row_data):
            row[i].text = str(v)

def add_rodape(doc, texto="BUENOSERV - SERVIÇOS DE ENGENHARIA LTDA"):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.text = texto
        p.style.font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def gerar_relatorio(nome, titulo, subtitulo, secoes, cliente="", data=""):
    doc = nova_pagina()
    add_capa(doc, titulo, subtitulo, cliente, data)
    add_rodape(doc)
    for secao in secoes:
        add_secao(doc, secao['titulo'], secao['conteudo'])
    salvar(doc, nome)

def gerar_ata(nome, projeto, data, participantes, pauta, decisoes, pendenciais):
    doc = nova_pagina()
    doc.add_heading("ATA DE REUNIÃO", level=0)
    for label, val in [("Projeto", projeto), ("Data", data), ("Participantes", participantes)]:
        p = doc.add_paragraph()
        r = p.add_run(f"{label}: ")
        r.bold = True
        p.add_run(val)
    doc.add_heading("Pauta", level=1)
    for item in pauta:
        doc.add_paragraph(item, style='List Number')
    doc.add_heading("Decisões", level=1)
    for d in decisoes:
        doc.add_paragraph(d, style='List Bullet')
    doc.add_heading("Pendências", level=1)
    add_tabela(doc, ["Item", "Responsável", "Prazo"], pendenciais)
    salvar(doc, nome)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: gen_docx.py <comando> <args_json>")
        print("Comandos: relatorio, ata")
        sys.exit(1)
    cmd = sys.argv[1]
    args = json.loads(sys.argv[2]) if len(sys.argv) > 2 else {}
    if cmd == "relatorio":
        gerar_relatorio(args["nome"], args["titulo"], args.get("subtitulo",""),
                       args["secoes"], args.get("cliente",""), args.get("data",""))
    elif cmd == "ata":
        gerar_ata(args["nome"], args["projeto"], args["data"], args["participantes"],
                 args["pauta"], args["decisoes"], args["pendenciais"])
    else:
        print(f"Comando desconhecido: {cmd}")
